#Python3.5
#2018/2/14
#参考教程：http://blog.csdn.net/c406495762/article/details/72331737#31-selenium
#待改进：中文字体；代码的通用性；无法爬取图片
 
from selenium import webdriver  #webdriver用来打开网页
from bs4 import BeautifulSoup   #用来爬取内容
import time     #用来等待完全加载
from docx import Document       #新建文档
from docx.enum.text import WD_ALIGN_PARAGRAPH   #用来居中显示标题    
 
def find_doc(driver, i):
    time.sleep(3)
    html = driver.page_source
    soup1 = BeautifulSoup(html, 'html.parser')
 
    result = soup1.find('div', attrs = {'class':'doc-title'} )
    doc_title = result.get_text()   ###得到文档标题
 
    try:
        elem = driver.find_element_by_xpath("//div[@data-flod-fun='continue-read']")
        elem.click()
        global doc_content_list
        doc_content_list = []
    except:
        pass
 
    result2 = soup1.find_all('p', attrs = {'class':'txt'} )
    for each in result2:
        text2 = each.get_text()
         
        if '            ' in text2:
            text3 = text2.replace( '            ', '' )
        else:
            text3 = text2
             
        doc_content_list.append(text3)  ###得到正文内容
         
    try:
        elem = driver.find_element_by_xpath("//div[@class='x-page next']")
        elem.click()
        print("已获取第 %d 页的内容。" % i)
        i += 1
        find_doc(driver, i)
    except:
        print("已获取所有共 %d 页的内容。\n" % i)
     
    return doc_title, doc_content_list
     
def save(doc_title, doc_content_list):
    document = Document()
    heading = document.add_heading(doc_title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
     
    for each in doc_content_list:
        paragraph = document.add_paragraph(each)
         
    document.save('百度文库-%s.docx' % doc_title)
    driver.quit()
    print("\n\n已全部写入文件    %s.docx,   请查收。" % doc_title)
         
if __name__ == "__main__":
    
    driver=webdriver.Chrome(executable_path="C:\Program Files (x86)\Google\Chrome\Application\chromedriver.exe")
 
    url = 'https://wenku.baidu.com/view/aa31a84bcf84b9d528ea7a2c.html'
    driver.get(url)
     
    m = 1
    title, content = find_doc(driver, m)
    save(title, content)
